from __future__ import annotations

import re
from typing import Any, Dict, List, Optional, Set, Tuple

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import Pt, RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_COLOR_INDEX

# Constants - if you keep a separate constants.py, ensure they match
PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_.-]+)\s*\}\}")
BOLD_MARKER_RE = re.compile(r"(\*\*|__)(.+?)\1", re.DOTALL)
STOPWORDS = {
    "and",
    "or",
    "the",
    "a",
    "an",
    "of",
    "for",
    "to",
    "in",
    "on",
    "with",
    "by",
    "&",
    "at",
    "as",
    "from",
}
# Detect top-level numbered headings like "4. Scope" or "5.1 Layout"
TOP_LEVEL_HEADING_RE = re.compile(r"^\d+(\.\d+)*\s+")
SECTION_HEADING_LINE_RE = re.compile(r"^\d+(?:\.\d+)+\s+")
SIMPLE_NUMBERED_LIST_RE = re.compile(r"^\d+[\.\)]\s+")
BULLET_PREFIX_CHARS = {"•", "-", "*", "–", "—", "·"}

# Run / Paragraph formatting helpers


def capture_run_formatting(paragraph: Paragraph) -> Dict[str, Any]:
    """
    Capture approximate run formatting from the first run of the paragraph.
    This is pragmatic: if a paragraph contains multiple inline formats, we preserve
    the first run's formatting as a default for newly inserted runs.
    """
    fmt: Dict[str, Any] = {}
    if not paragraph.runs:
        return fmt
    run = paragraph.runs[0]
    font = run.font
    if run.bold is not None:
        fmt["bold"] = bool(run.bold)
    if run.italic is not None:
        fmt["italic"] = bool(run.italic)
    if run.underline is not None:
        fmt["underline"] = run.underline
    if font is not None:
        if font.name:
            fmt["font_name"] = font.name
        if font.size:
            try:
                fmt["font_size_pt"] = float(font.size.pt)
            except Exception:
                pass
        color = getattr(font, "color", None)
        if color is not None:
            rgb = getattr(color, "rgb", None)
            if rgb is not None:
                try:
                    fmt["font_color_rgb"] = rgb.hex
                except Exception:
                    fmt["font_color_rgb"] = str(rgb)
            theme = getattr(color, "theme_color", None)
            if theme is not None:
                fmt["font_color_theme"] = getattr(theme, "name", str(theme))
        highlight = getattr(font, "highlight_color", None)
        if highlight is not None:
            fmt["highlight_color"] = getattr(highlight, "name", str(highlight))
    return fmt


def apply_run_formatting(run: Run, formatting: Dict[str, Any] | None) -> None:
    if not formatting:
        return
    if "bold" in formatting:
        run.bold = formatting.get("bold")
    if "italic" in formatting:
        run.italic = formatting.get("italic")
    if "underline" in formatting:
        run.underline = formatting.get("underline")
    font = run.font
    if font is None:
        return
    fn = formatting.get("font_name")
    if fn:
        try:
            font.name = fn
        except Exception:
            pass
    fs = formatting.get("font_size_pt")
    if fs is not None:
        try:
            font.size = Pt(float(fs))
        except Exception:
            pass
    rgb = formatting.get("font_color_rgb")
    if rgb:
        try:
            font.color.rgb = RGBColor.from_string(rgb)
        except Exception:
            pass
    theme = formatting.get("font_color_theme")
    if theme:
        enum_value = getattr(MSO_THEME_COLOR_INDEX, str(theme), None)
        if enum_value is not None:
            try:
                font.color.theme_color = enum_value
            except Exception:
                pass
    highlight = formatting.get("highlight_color")
    if highlight:
        enum_value = getattr(WD_COLOR_INDEX, str(highlight), None)
        if enum_value is not None:
            try:
                font.highlight_color = enum_value
            except Exception:
                pass


def capture_paragraph_properties(paragraph: Paragraph) -> Optional[str]:
    """Return the raw <w:pPr> XML if present, else None."""
    try:
        ppr = paragraph._p.pPr
        return ppr.xml if ppr is not None else None
    except Exception:
        return None


def capture_list_properties(paragraph: Paragraph) -> Optional[str]:
    """
    Return the raw <w:numPr> xml for the paragraph if present.
    This is the best way to reliably preserve bullets/numbering.
    """
    try:
        ppr = paragraph._p.pPr
        if ppr is None:
            return None
        numpr = getattr(ppr, "numPr", None)
        return numpr.xml if numpr is not None else None
    except Exception:
        return None


def apply_paragraph_properties(paragraph: Paragraph, ppr_xml: Optional[str]) -> None:
    """Apply raw <w:pPr> xml to paragraph._p (replacing existing pPr when possible)."""
    if not ppr_xml:
        return
    try:
        new_ppr = parse_xml(ppr_xml)
    except Exception:
        return
    p = paragraph._p
    existing = getattr(p, "pPr", None)
    if existing is not None:
        try:
            p.remove(existing)
        except Exception:
            pass
    try:
        p.insert(0, new_ppr)
    except Exception:
        try:
            p.append(new_ppr)
        except Exception:
            pass


def apply_list_properties(paragraph: Paragraph, numpr_xml: Optional[str]) -> None:
    """Apply raw <w:numPr> xml into paragraph pPr (creating pPr if needed)."""
    if not numpr_xml:
        return
    try:
        new_numpr = parse_xml(numpr_xml)
    except Exception:
        return
    p = paragraph._p
    pPr = getattr(p, "pPr", None)
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    existing = getattr(pPr, "numPr", None)
    if existing is not None:
        try:
            pPr.remove(existing)
        except Exception:
            pass
    try:
        pPr.append(new_numpr)
    except Exception:
        pass


# Paragraph insertion / replacement helpers


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    run_format: Dict[str, Any] | None = None,
    ppr_xml: Optional[str] = None,
    numpr_xml: Optional[str] = None,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_par = Paragraph(new_p, paragraph._parent)
    if ppr_xml:
        apply_paragraph_properties(new_par, ppr_xml)
    if numpr_xml:
        apply_list_properties(new_par, numpr_xml)
    if text:
        r = new_par.add_run(text)
        apply_run_formatting(r, run_format)
    return new_par


def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)


def _markdown_segments(text: str) -> List[Dict[str, Any]]:
    segments: List[Dict[str, Any]] = []
    last = 0
    for match in BOLD_MARKER_RE.finditer(text):
        start, end = match.span()
        if start > last:
            segments.append({"text": text[last:start], "bold": False})
        segments.append({"text": match.group(2), "bold": True})
        last = end
    if last < len(text):
        segments.append({"text": text[last:], "bold": False})
    if not segments:
        segments.append({"text": text, "bold": False})
    return segments


def _apply_markdown_runs(paragraph: Paragraph, text: str, run_format: Dict[str, Any] | None) -> None:
    for segment in _markdown_segments(text):
        if not segment["text"]:
            continue
        run = paragraph.add_run(segment["text"])
        apply_run_formatting(run, run_format)
        if segment.get("bold"):
            run.bold = True


def replace_paragraph_text(
    paragraph: Paragraph,
    text: str,
    run_format: Dict[str, Any] | None = None,
    ppr_xml: Optional[str] = None,
    numpr_xml: Optional[str] = None,
) -> None:
    """
    Replace paragraph children except pPr, then add a single run with `text`.
    Reapply captured paragraph properties and numPr (list) so bullets/nesting persist.
    """
    if run_format is None:
        run_format = capture_run_formatting(paragraph)
    if ppr_xml is None:
        ppr_xml = capture_paragraph_properties(paragraph)
    if numpr_xml is None:
        numpr_xml = capture_list_properties(paragraph)

    p_el = paragraph._p
    # remove all children except pPr
    for child in list(p_el):
        if child.tag.endswith("}pPr"):
            continue
        p_el.remove(child)

    if text:
        _apply_markdown_runs(paragraph, text, run_format)

    if ppr_xml:
        apply_paragraph_properties(paragraph, ppr_xml)
    if numpr_xml:
        apply_list_properties(paragraph, numpr_xml)


# Heuristic heading detection / section inference (improved)


def looks_like_heading(paragraph: Paragraph) -> bool:
    """
    Improved heuristics for heading detection:
    - explicit Heading style names (Heading, Title, Subtitle)
    - numbered headings (4.1, 5.2)
    - bold with short length
    - larger font size relative to neighbors
    - all-caps short lines
    - lines that end with ':' often are headings
    """
    text = paragraph.text.strip()
    if not text:
        return False

    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    if style_name and any(token.lower() in style_name.lower() for token in ("heading", "title", "subtitle")):
        return True

    if TOP_LEVEL_HEADING_RE.match(text):
        return True

    # avoid lists
    if text.startswith(("•", "-", "*", "–", "—", "·")):
        return False

    # short strongly styled paragraphs
    if len(text) <= 120:
        # check runs for bold majority or first run bold and short
        try:
            runs = paragraph.runs
            if runs:
                bold_runs = sum(1 for r in runs if getattr(r, "bold", None))
                if bold_runs and bold_runs >= max(1, len(runs) // 2):
                    return True
                # first run bold and short heading-like
                if getattr(runs[0], "bold", None) and len(text.split()) <= 8:
                    return True
        except Exception:
            pass

    # check font size heuristics: if first run font size is larger than typical (>= 14pt) consider heading
    try:
        runs = paragraph.runs
        if runs:
            first_font = runs[0].font
            size = getattr(first_font, "size", None)
            if size and getattr(size, "pt", None) and size.pt >= 13.5:
                return True
    except Exception:
        pass

    # ALL CAPS short lines
    alpha_chars = [ch for ch in text if ch.isalpha()]
    if alpha_chars and sum(1 for ch in alpha_chars if ch.isupper()) / len(alpha_chars) > 0.7 and len(text.split()) <= 6:
        return True

    # endswith colon, often heading-like
    if text.endswith(":") and len(text.split()) <= 8:
        return True

    # fallback false
    return False


def slugify(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "_", text.strip().lower()).strip("_") or "section"


def infer_sections(doc: Document) -> List[Dict[str, Any]]:
    """
    Identify sections by headings; collect paragraph metadata (style/pPr/numPr/run_format)
    so that replacements preserve formatting and lists.
    """
    sections: List[Dict[str, Any]] = []
    seen_ids: Set[str] = set()
    numbered_mode = any(is_top_level_numbered_heading(p) for p in doc.paragraphs)

    for idx, paragraph in enumerate(doc.paragraphs):
        if numbered_mode:
            if not is_top_level_numbered_heading(paragraph):
                continue
        else:
            if not looks_like_heading(paragraph):
                continue

        heading_text = paragraph.text.strip()
        candidate = slugify(heading_text)
        unique = candidate
        counter = 2
        while unique in seen_ids:
            unique = f"{candidate}_{counter}"
            counter += 1
        seen_ids.add(unique)
        sections.append({"id": unique, "heading": heading_text, "heading_index": idx})

    if not sections:
        return []

    total_paragraphs = len(doc.paragraphs)
    for i, section in enumerate(sections):
        start = section["heading_index"] + 1
        end = sections[i + 1]["heading_index"] if i + 1 < len(sections) else total_paragraphs

        # skip leading/trailing blank paragraphs inside section
        while start < end and not doc.paragraphs[start].text.strip():
            start += 1
        while end > start and not doc.paragraphs[end - 1].text.strip():
            end -= 1

        section["content_start"] = start
        section["content_end"] = end

        # collect paragraph metadata for reapplication
        paragraph_meta: List[Dict[str, Any]] = []
        for pi in range(start, end):
            p = doc.paragraphs[pi]
            style_name = getattr(getattr(p, "style", None), "name", None)
            ppr_xml = capture_paragraph_properties(p)
            numpr_xml = capture_list_properties(p)
            run_fmt = capture_run_formatting(p)
            paragraph_meta.append(
                {
                    "index": pi,
                    "style": style_name,
                    "ppr_xml": ppr_xml,
                    "numpr_xml": numpr_xml,
                    "run_format": run_fmt,
                }
            )

        section["paragraphs"] = paragraph_meta
        section["default_text"] = "\n\n".join(
            doc.paragraphs[j].text.strip() for j in range(start, end) if doc.paragraphs[j].text.strip()
        )

    return sections


def is_top_level_numbered_heading(paragraph: Paragraph) -> bool:
    return bool(TOP_LEVEL_HEADING_RE.match(paragraph.text.strip() or ""))


# Tables summarizer


def clone_table_rows(rows: List[List[str]]) -> List[List[str]]:
    return [[("" if cell is None else str(cell)) for cell in row] for row in rows]


def friendly_label_from_name(name: str) -> str:
    cleaned = name.replace("_", " ").replace("-", " ").replace(".", " ")
    parts = [part.strip() for part in cleaned.split() if part.strip()]
    if not parts:
        return ""
    return " ".join(part.capitalize() for part in parts)


def summarize_tables(doc: Document, sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    paragraph_to_index = {paragraph._p: idx for idx, paragraph in enumerate(doc.paragraphs)}
    section_by_heading = {section.get("heading_index"): section.get("id") for section in sections}
    summary: List[Dict[str, Any]] = []

    for table_index, table in enumerate(doc.tables):
        table_rows: List[List[str]] = []
        row_meta: List[List[Dict[str, int]]] = []
        for row_idx, row in enumerate(table.rows):
            values: List[str] = []
            meta_cells: List[Dict[str, int]] = []
            for col_idx, cell in enumerate(row.cells):
                values.append(cell.text.strip())
                meta_cells.append({"row": row_idx, "col": col_idx})
            table_rows.append(values)
            row_meta.append(meta_cells)

        # locate nearest preceding paragraph to associate table with a section
        section_id = None
        previous = table._tbl.getprevious()
        while previous is not None:
            paragraph_index = paragraph_to_index.get(previous)
            if paragraph_index is not None:
                for section in sections:
                    start = section.get("content_start")
                    end = section.get("content_end")
                    if start is None or end is None:
                        continue
                    if start <= paragraph_index <= max(start, end - 1):
                        section_id = section.get("id")
                        break
                if section_id is None and paragraph_index in section_by_heading:
                    section_id = section_by_heading[paragraph_index]
                if section_id is not None:
                    break
            previous = previous.getprevious()

        first_cell = ""
        if table_rows and table_rows[0]:
            first_cell = table_rows[0][0].strip()

        summary.append(
            {
                "table_index": table_index,
                "rows": clone_table_rows(table_rows),
                "cell_meta": row_meta,
                "section_id": section_id,
                "label": first_cell or f"Table {table_index + 1}",
            }
        )

    return summary


def build_schema(
    doc: Document,
    sections: List[Dict[str, Any]],
    placeholders: List[str],
    tables: List[Dict[str, Any]],
):
    schema: List[Dict[str, Any]] = []
    field_map: Dict[str, Any] = {}
    tables_by_section: Dict[str | None, List[Dict[str, Any]]] = {}
    applied_tables: Set[int] = set()

    for table in tables:
        section_key = table.get("section_id")
        tables_by_section.setdefault(section_key, []).append(table)

    def attach_table_field(target_fields: List[Dict[str, Any]], table_info: Dict[str, Any]):
        prefix = table_info.get("section_id")
        if prefix:
            field_name = f"{prefix}_table_{table_info['table_index'] + 1}"
        else:
            field_name = f"table_{table_info['table_index'] + 1}"
        table_label = table_info.get("label") or friendly_label_from_name(field_name)

        field_map[field_name] = {
            "type": "table",
            "table_index": table_info["table_index"],
            "rows": table_info.get("cell_meta", []),
        }

        applied_tables.add(table_info["table_index"])

        target_fields.append(
            {
                "name": field_name,
                "type": "table",
                "label": table_label,
                "rows": clone_table_rows(table_info.get("rows", [])),
                "default": clone_table_rows(table_info.get("rows", [])),
            }
        )

    for section in sections:
        fields_public: List[Dict[str, Any]] = []

        field_name = f"{section['id']}_content"
        field_map[field_name] = {
            "type": "section",
            "section": section.copy(),
        }

        fields_public.append(
            {
                "name": field_name,
                "type": "textarea",
                "label": section.get("heading", friendly_label_from_name(field_name)),
                "default": section.get("default_text", ""),
                "polishable": True,
            }
        )

        for table_info in tables_by_section.get(section.get("id"), []):
            attach_table_field(fields_public, table_info)

        schema.append(
            {
                "section": section.get("heading", ""),
                "fields": fields_public,
            }
        )

    if not sections:
        schema = []

    for table in tables:
        if table["table_index"] in applied_tables:
            continue

        field_block: List[Dict[str, Any]] = []
        attach_table_field(field_block, table)
        schema.append(
            {
                "section": table.get("label") or friendly_label_from_name(field_block[0]["name"]),
                "fields": field_block,
            }
        )

    return schema, field_map


# List detection (less aggressive, uses numPr + style heuristics)


def paragraph_has_numpr(paragraph: Paragraph) -> bool:
    """Return True if paragraph has an internal numPr (numbering/bullet properties)."""
    try:
        ppr = paragraph._p.pPr
        if ppr is None:
            return False
        return getattr(ppr, "numPr", None) is not None
    except Exception:
        return False


def paragraph_is_list(paragraph: Paragraph) -> bool:
    """
    Consider a paragraph a list item if:
    - it has explicit numPr, OR
    - it begins with a bullet glyph AND not obviously a heading.
    We avoid classifying bold short headings as list items.
    """
    if paragraph_has_numpr(paragraph):
        return True

    text = paragraph.text.strip()
    if not text:
        return False

    # bullets starting glyphs - still require not-a-heading
    if text[0] in {"•", "-", "*", "–", "—", "·"}:
        # if the paragraph looks like a heading, don't treat as list
        if looks_like_heading(paragraph):
            return False
        return True

    # also check if style explicitly named list
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    if style_name and any(tok in style_name.lower() for tok in ("list", "bullet", "number")):
        if not looks_like_heading(paragraph):
            return True

    return False


def clean_bullet_text(text: str) -> str:
    text = text.strip()
    # remove leading bullet glyphs / numbering characters
    text = re.sub(r"^[\u2022\u2023\u25E6\u2043\u2219\-\*\u2013\u2014\.\)\s]+", "", text)
    return text.strip()


def text_looks_like_list_item(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if stripped[0] in BULLET_PREFIX_CHARS:
        return True
    if SIMPLE_NUMBERED_LIST_RE.match(stripped):
        return True
    return False


def _line_is_section_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped[0] in BULLET_PREFIX_CHARS:
        return False
    if SECTION_HEADING_LINE_RE.match(stripped):
        return True
    words = [word for word in re.split(r"\s+", stripped) if word]
    alpha_words = [word for word in words if word[0].isalpha()]
    if 2 <= len(alpha_words) <= 8 and all(word[0].isupper() for word in alpha_words):
        return True
    return False


def _ensure_section_breaks(text: str) -> str:
    lines = text.split("\n")
    rebuilt: List[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped and _line_is_section_heading(stripped) and rebuilt and rebuilt[-1].strip() != "":
            rebuilt.append("")
        rebuilt.append(line)
    return "\n".join(rebuilt)


# Field application


def set_cell_text(cell, text: Any) -> None:
    value = "" if text is None else str(text)
    if not cell.paragraphs:
        cell.add_paragraph("")
    primary = cell.paragraphs[0]
    replace_paragraph_text(primary, value)
    for extra in cell.paragraphs[1:]:
        remove_paragraph(extra)


def delete_table_row(table, index: int) -> None:
    if index < 0 or index >= len(table.rows):
        return
    table._tbl.remove(table.rows[index]._tr)


def apply_table_field(doc: Document, field_meta: Dict[str, Any], value: Any) -> None:
    table_index = field_meta.get("table_index")
    if table_index is None or table_index < 0 or table_index >= len(doc.tables):
        return

    table = doc.tables[table_index]
    value_rows: List[List[Any]] = []
    if isinstance(value, list):
        for row in value:
            if isinstance(row, list):
                value_rows.append([(cell if cell is not None else "") for cell in row])
            elif isinstance(row, dict):
                normalized_row = []
                for col_idx in range(len(table.columns)):
                    normalized_row.append(row.get(str(col_idx), ""))
                value_rows.append(normalized_row)
            else:
                value_rows.append([row])
    else:
        value_rows.append([value])

    target_rows = max(len(value_rows), 1)

    while len(table.rows) < target_rows:
        table.add_row()

    while len(table.rows) > target_rows:
        delete_table_row(table, len(table.rows) - 1)

    max_columns = max((len(row.cells) for row in table.rows), default=0)

    for r_idx, row in enumerate(table.rows):
        row_values: List[Any]
        if r_idx < len(value_rows):
            row_values = value_rows[r_idx]
        else:
            row_values = [""] * max_columns
        for c_idx, cell in enumerate(row.cells):
            cell_value = row_values[c_idx] if c_idx < len(row_values) else ""
            set_cell_text(cell, cell_value)


def apply_checkbox_field(doc: Document, field_meta: Dict[str, Any], value: Any) -> None:
    entries = {}
    if isinstance(value, list):
        for entry in value:
            if isinstance(entry, dict) and "value" in entry:
                entries[entry["value"]] = {
                    "selected": bool(entry.get("selected", False)),
                    "text": entry.get("text", ""),
                }
            elif isinstance(entry, str):
                entries[entry] = {"selected": True, "text": ""}

    options = sorted(
        field_meta.get("options", []), key=lambda opt: opt.get("paragraph_index", 0), reverse=True
    )

    for option in options:
        idx = int(option.get("paragraph_index", 0))
        if idx >= len(doc.paragraphs) or idx < 0:
            continue

        paragraph = doc.paragraphs[idx]
        style = option.get("style")
        if style:
            try:
                paragraph.style = style
            except Exception:
                pass

        entry = entries.get(option["value"]) or {"selected": True, "text": ""}
        if not entry.get("selected", False):
            remove_paragraph(paragraph)
            continue

        new_text = entry.get("text") or option.get("text", "")
        replace_paragraph_text(
            paragraph,
            new_text,
            capture_run_formatting(paragraph),
            capture_paragraph_properties(paragraph),
            capture_list_properties(paragraph),
        )


def split_paragraph_blocks(text: str, target_count: Optional[int] = None) -> List[str]:
    """
    Split text into paragraph blocks separated by one or more blank lines.
    If target_count provided, try to match that count: either split lines or pad/truncate.
    This function tries to avoid duplicating original paragraphs by returning exactly target_count blocks.
    """
    if text is None:
        return [""]

    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = _ensure_section_breaks(normalized)
    if not normalized.strip():
        blocks = [""]
    else:
        # primary split on two-or-more consecutive newlines
        blocks = [block.strip() for block in re.split(r"\n{2,}", normalized) if block is not None]
        if target_count and len(blocks) < target_count:
            # fallback: try single-line splits; keep all lines so new bullets aren't truncated
            lines = [ln.strip() for ln in normalized.split("\n") if ln.strip()]
            if len(lines) >= target_count:
                blocks = lines
        # final cleanup
        blocks = [b for b in blocks]
    # pad or truncate to target_count exactly if provided
    if target_count and len(blocks) < target_count:
        blocks = blocks + [""] * (target_count - len(blocks))
    return blocks or [""]


def apply_textarea_field(doc: Document, field_meta: Dict[str, Any], value: Any) -> None:
    """
    Replace a paragraph block region (start..end) with value, preserving paragraph-level
    metadata (run_format, pPr, numPr) for each inner paragraph. Avoid duplication:
    - use paragraph_metas to control exact inserted paragraph count.
    """

    start = int(field_meta.get("start", 0))
    end = int(field_meta.get("end", start))
    style = field_meta.get("style")
    heading_index = field_meta.get("heading_index")
    paragraph_metas = field_meta.get("paragraphs") or []

    def paragraph_meta_for(index: int, block_text: Optional[str] = None) -> Dict[str, Any]:
        if paragraph_metas:
            if index < len(paragraph_metas):
                base = paragraph_metas[index]
            else:
                base = paragraph_metas[-1]
            meta = dict(base)
        else:
            meta = {}
        if block_text is not None and not text_looks_like_list_item(block_text):
            meta["numpr_xml"] = None
        return meta

    if not doc.paragraphs:
        doc.add_paragraph("")

    # ensure start/end are within bounds; if start beyond doc, insert after heading if present
    if start >= len(doc.paragraphs):
        if heading_index is not None and 0 <= heading_index < len(doc.paragraphs):
            heading_para = doc.paragraphs[heading_index]
            new_para = insert_paragraph_after(heading_para, "")
            paragraphs = list(doc.paragraphs)
            start = paragraphs.index(new_para)
            end = start + 1
        else:
            doc.add_paragraph("")
            start = len(doc.paragraphs) - 1
            end = start + 1
    else:
        start = max(0, start)
        end = max(start + 1, end)

    end = min(len(doc.paragraphs), max(start + 1, end))

    text_value = "" if value is None else str(value)
    target_count = len(paragraph_metas) if paragraph_metas else None
    blocks = split_paragraph_blocks(text_value, target_count)
    block_metas = [paragraph_meta_for(idx, block) for idx, block in enumerate(blocks)]

    normalized_blocks: List[str] = []
    for block, meta in zip(blocks, block_metas):
        text = block
        if meta.get("numpr_xml"):
            text = clean_bullet_text(text)
        normalized_blocks.append(text)
    blocks = normalized_blocks

    # ensure we don't accidentally duplicate the first original paragraph content.
    primary = doc.paragraphs[start]
    primary_meta = block_metas[0] if block_metas else {}
    run_fmt = primary_meta.get("run_format")
    ppr_xml = primary_meta.get("ppr_xml")
    numpr_xml = primary_meta.get("numpr_xml")

    # If block[0] equals original paragraph text and there is only one block and target_count==1,
    # then keep as-is (no replacement) to avoid duplication.
    original_primary_text = primary.text.strip()
    block0 = blocks[0].strip() if blocks else ""
    if target_count in (None, 1) and block0 and block0 == original_primary_text:
        # still may need to update formatting; reapply captured ppr/numpr if needed
        replace_paragraph_text(primary, original_primary_text, run_fmt, ppr_xml, numpr_xml)
    else:
        replace_paragraph_text(primary, blocks[0], run_fmt, ppr_xml, numpr_xml)

    if primary_meta.get("style"):
        try:
            primary.style = primary_meta.get("style")
        except Exception:
            pass
    if primary_meta.get("ppr_xml"):
        apply_paragraph_properties(primary, primary_meta.get("ppr_xml"))
    if primary_meta.get("numpr_xml"):
        apply_list_properties(primary, primary_meta.get("numpr_xml"))

    # remove leftover paragraphs in original range (reverse order to be safe)
    for idx in range(end - 1, start, -1):
        if idx < len(doc.paragraphs):
            remove_paragraph(doc.paragraphs[idx])

    # insert remaining blocks, matching paragraph_metas where available
    current = primary
    for bi, block in enumerate(blocks[1:], start=1):
        meta = block_metas[bi] if bi < len(block_metas) else paragraph_meta_for(bi)
        run_fmt = meta.get("run_format")
        ppr_xml = meta.get("ppr_xml")
        numpr_xml = meta.get("numpr_xml")

        current = insert_paragraph_after(current, block, run_fmt, ppr_xml, numpr_xml)
        if meta.get("style"):
            try:
                current.style = meta.get("style")
            except Exception:
                pass


def replace_section_content(doc: Document, section: Dict[str, Any], value: Any) -> None:
    text_value = "" if value is None else str(value)
    meta = {
        "start": section.get("content_start", 0),
        "end": section.get("content_end", 0),
        "heading_index": section.get("heading_index"),
        "paragraphs": section.get("paragraphs", []),
    }
    apply_textarea_field(doc, meta, text_value)


def apply_field(doc: Document, field_meta: Dict[str, Any], value: Any) -> None:
    ftype = field_meta.get("type")
    if ftype in {"text", "textarea", "date"}:
        apply_textarea_field(doc, field_meta, value)
    elif ftype == "checkbox-group":
        apply_checkbox_field(doc, field_meta, value)
    elif ftype == "table":
        apply_table_field(doc, field_meta, value)
    elif ftype == "section":
        replace_section_content(doc, field_meta.get("section", {}), value)


# Placeholder extraction + replacement


def extract_placeholders(doc: Document) -> List[str]:
    found: Set[str] = set()
    for p in doc.paragraphs:
        for m in PLACEHOLDER_RE.findall(p.text):
            found.add(m)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for m in PLACEHOLDER_RE.findall(p.text):
                        found.add(m)
    try:
        for s in doc.sections:
            for p in s.header.paragraphs:
                for m in PLACEHOLDER_RE.findall(p.text):
                    found.add(m)
            for p in s.footer.paragraphs:
                for m in PLACEHOLDER_RE.findall(p.text):
                    found.add(m)
    except Exception:
        pass
    return sorted(found)


def replace_placeholders(doc: Document, data: Dict[str, Any], metadata: Optional[Dict[str, Any]] = None) -> Document:
    """
    Replace inline placeholders for scalar values while preserving formatting.
    After that, apply structured fields (field_map) if metadata is provided.
    """

    def replace_in_paragraph(paragraph: Paragraph):
        orig_text = paragraph.text
        new_text = orig_text
        for key, value in data.items():
            # skip structured fields (list/table/section) until later
            if isinstance(value, (dict, list)):
                continue
            pattern = re.compile(rf"\{{\{{\s*{re.escape(str(key))}\s*\}}\}}")
            replacement = "" if value is None else str(value)
            new_text = pattern.sub(replacement, new_text)
        if new_text != orig_text:
            # capture format and reapply
            run_fmt = capture_run_formatting(paragraph)
            ppr = capture_paragraph_properties(paragraph)
            numpr = capture_list_properties(paragraph)
            replace_paragraph_text(paragraph, new_text, run_fmt, ppr, numpr)

    # body
    for p in list(doc.paragraphs):
        replace_in_paragraph(p)

    # tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    replace_in_paragraph(p)

    # headers/footers
    try:
        for s in doc.sections:
            for p in list(s.header.paragraphs):
                replace_in_paragraph(p)
            for p in list(s.footer.paragraphs):
                replace_in_paragraph(p)
    except Exception:
        pass

    # structured fields via metadata
    if not metadata:
        return doc

    field_map = metadata.get("field_map") or {}
    if field_map:
        ordered_fields = sorted(field_map.items(), key=lambda item: item[1].get("start", 0), reverse=True)
        for field_name, field_meta in ordered_fields:
            if field_name not in data:
                continue
            apply_field(doc, field_meta, data[field_name])
    else:
        # fallback: section-level replacement
        section_map = {section["id"]: section for section in metadata.get("auto_sections", [])}
        for key, value in data.items():
            section = section_map.get(key)
            if section:
                replace_section_content(doc, section, value)

    return doc


# Local test helper (optional)


def _test_local(input_path: str, output_path: str, payload: Dict[str, Any]) -> None:
    """
    Simple local test runner (not used by server) to read a docx, infer sections/tables,
    apply payload and save resulting docx.
    """
    doc = Document(input_path)
    placeholders = extract_placeholders(doc)
    sections = infer_sections(doc)
    tables = summarize_tables(doc, sections)
    metadata = {"placeholders": placeholders, "auto_sections": sections, "tables": tables}
    replace_placeholders(doc, payload, metadata)
    doc.save(output_path)


# End of file

